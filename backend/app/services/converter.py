"""文件格式转换服务"""
import os
import uuid
import subprocess
from datetime import datetime
from typing import Optional
from PIL import Image
from sqlalchemy.orm import Session

from ..models.file import File
from ..models.conversion import Conversion
from ..config import settings
from ..utils.file_utils import get_file_extension, ensure_directory_exists


class ConverterService:
    """文件格式转换服务类"""

    def __init__(self, db: Session):
        """
        初始化转换服务

        Args:
            db: 数据库会话
        """
        self.db = db
        ensure_directory_exists(settings.OUTPUT_DIR)

    def convert_image_to_pdf(self, file_id: int) -> Conversion:
        """
        将图片转换为 PDF

        Args:
            file_id: 文件ID

        Returns:
            Conversion: 转换任务记录

        Raises:
            ValueError: 文件不存在或不是图片文件
            Exception: 转换过程中的错误
        """
        # 获取文件记录
        db_file = self.db.query(File).filter(File.id == file_id).first()
        if not db_file:
            raise ValueError(f"文件不存在: {file_id}")

        # 检查文件类型
        file_ext = get_file_extension(db_file.original_name)
        if file_ext not in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            raise ValueError(f"不支持的图片格式: {file_ext}")

        # 检查源文件是否存在
        if not os.path.exists(db_file.file_path):
            raise ValueError(f"源文件不存在: {db_file.file_path}")

        # 创建转换任务记录
        conversion = Conversion(
            file_id=file_id,
            source_format=file_ext,
            target_format='pdf',
            status='processing'
        )
        self.db.add(conversion)
        self.db.commit()
        self.db.refresh(conversion)

        try:
            # 生成输出文件名和路径
            output_filename = f"{uuid.uuid4()}.pdf"
            output_path = os.path.join(settings.OUTPUT_DIR, output_filename)

            # 使用 Pillow 进行转换
            image = Image.open(db_file.file_path)

            # 如果图片是 RGBA 模式，转换为 RGB（PDF 不支持透明度）
            if image.mode in ('RGBA', 'LA', 'P'):
                # 创建白色背景
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
                image = background
            elif image.mode != 'RGB':
                image = image.convert('RGB')

            # 保存为 PDF
            image.save(output_path, 'PDF', resolution=100.0, quality=95)

            # 更新转换任务状态
            conversion.status = 'completed'
            conversion.result_path = output_path
            conversion.result_filename = output_filename
            conversion.completed_at = datetime.now()

            # 更新原文件状态
            db_file.status = 'converted'

            self.db.commit()
            self.db.refresh(conversion)

            return conversion

        except Exception as e:
            # 转换失败，更新状态
            conversion.status = 'failed'
            conversion.error_message = str(e)
            conversion.completed_at = datetime.now()
            self.db.commit()
            raise Exception(f"图片转 PDF 失败: {str(e)}")

    def convert_word_to_pdf(self, file_id: int) -> Conversion:
        """
        将 Word 文档转换为 PDF

        Args:
            file_id: 文件ID

        Returns:
            Conversion: 转换任务记录

        Raises:
            ValueError: 文件不存在或不是 Word 文件
            Exception: 转换过程中的错误
        """
        # 获取文件记录
        db_file = self.db.query(File).filter(File.id == file_id).first()
        if not db_file:
            raise ValueError(f"文件不存在: {file_id}")

        # 检查文件类型
        file_ext = get_file_extension(db_file.original_name)
        if file_ext not in ['doc', 'docx']:
            raise ValueError(f"不支持的 Word 格式: {file_ext}")

        # 检查源文件是否存在
        if not os.path.exists(db_file.file_path):
            raise ValueError(f"源文件不存在: {db_file.file_path}")

        # 创建转换任务记录
        conversion = Conversion(
            file_id=file_id,
            source_format=file_ext,
            target_format='pdf',
            status='processing'
        )
        self.db.add(conversion)
        self.db.commit()
        self.db.refresh(conversion)

        try:
            # 生成输出文件名和路径
            output_filename = f"{uuid.uuid4()}.pdf"
            output_path = os.path.join(settings.OUTPUT_DIR, output_filename)

            # 尝试使用 LibreOffice 转换（如果可用）
            success = self._convert_with_libreoffice(db_file.file_path, output_path)

            if not success:
                # 如果 LibreOffice 不可用，使用备用方案
                success = self._convert_word_fallback(db_file.file_path, output_path)

            if not success:
                raise Exception("Word 转 PDF 失败：未找到可用的转换工具")

            # 更新转换任务状态
            conversion.status = 'completed'
            conversion.result_path = output_path
            conversion.result_filename = output_filename
            conversion.completed_at = datetime.now()

            # 更新原文件状态
            db_file.status = 'converted'

            self.db.commit()
            self.db.refresh(conversion)

            return conversion

        except Exception as e:
            # 转换失败，更新状态
            conversion.status = 'failed'
            conversion.error_message = str(e)
            conversion.completed_at = datetime.now()
            self.db.commit()
            raise Exception(f"Word 转 PDF 失败: {str(e)}")

    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> bool:
        """
        使用 LibreOffice 转换 Word 到 PDF

        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径

        Returns:
            是否转换成功
        """
        try:
            # 尝试查找 LibreOffice
            libreoffice_paths = [
                "libreoffice",  # Linux/Mac
                "soffice",
                r"C:\Program Files\LibreOffice\program\soffice.exe",  # Windows
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]

            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run(
                        [path, "--version"],
                        capture_output=True,
                        timeout=5
                    )
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue

            if not libreoffice_cmd:
                return False

            # 使用 LibreOffice 转换
            output_dir = os.path.dirname(output_path)
            result = subprocess.run(
                [
                    libreoffice_cmd,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", output_dir,
                    input_path
                ],
                capture_output=True,
                timeout=60,
                text=True
            )

            if result.returncode != 0:
                print(f"LibreOffice 转换失败: {result.stderr}")
                return False

            # LibreOffice 会生成一个临时文件名，需要重命名
            input_basename = os.path.splitext(os.path.basename(input_path))[0]
            temp_output = os.path.join(output_dir, f"{input_basename}.pdf")

            if os.path.exists(temp_output):
                os.rename(temp_output, output_path)
                return True

            return False

        except Exception as e:
            print(f"LibreOffice 转换异常: {e}")
            return False

    def _convert_word_fallback(self, input_path: str, output_path: str) -> bool:
        """
        Word 转 PDF 备用方案（使用 python-docx + reportlab）

        注意：此方法仅支持基本文本和表格，不支持图片和复杂格式
        建议安装 LibreOffice 以获得完整的转换支持

        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径

        Returns:
            是否转换成功
        """
        try:
            from docx import Document
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import io

            # 读取 Word 文档
            doc = Document(input_path)

            # 创建 PDF
            pdf = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )

            # 注册中文字体
            try:
                # 尝试使用 Windows 系统字体
                font_paths = [
                    'C:/Windows/Fonts/simsun.ttc',  # 宋体
                    'C:/Windows/Fonts/simhei.ttf',  # 黑体
                    'C:/Windows/Fonts/msyh.ttc',    # 微软雅黑
                ]
                chinese_font = None
                for font_path in font_paths:
                    if os.path.exists(font_path):
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            chinese_font = 'ChineseFont'
                            print(f"成功注册中文字体: {font_path}")
                            break
                        except Exception as e:
                            print(f"注册字体失败 {font_path}: {e}")
                            continue

                if not chinese_font:
                    print("警告：未找到中文字体，中文可能无法正常显示")
                    chinese_font = 'Helvetica'
            except Exception as e:
                print(f"字体注册异常: {e}")
                chinese_font = 'Helvetica'

            styles = getSampleStyleSheet()

            # 创建支持中文的样式
            if chinese_font != 'Helvetica':
                chinese_style = ParagraphStyle(
                    name='ChineseStyle',
                    parent=styles['Normal'],
                    fontName=chinese_font,
                    fontSize=10,
                    leading=14,
                )
                styles.add(chinese_style)
            else:
                chinese_style = styles['Normal']

            story = []

            # 处理段落
            print(f"处理 {len(doc.paragraphs)} 个段落")
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    try:
                        # 使用 HTML 转义来处理特殊字符
                        from xml.sax.saxutils import escape
                        text = escape(text)
                        p = Paragraph(text, chinese_style)
                        story.append(p)
                        story.append(Spacer(1, 0.2*inch))
                    except Exception as e:
                        print(f"段落处理错误: {e}, 文本: {text[:50]}")
                        # 如果出错，添加纯文本
                        story.append(Paragraph(f"[文本: {text[:100]}...]", chinese_style))

            # 处理表格
            print(f"处理 {len(doc.tables)} 个表格")
            for idx, table in enumerate(doc.tables):
                try:
                    # 提取表格数据
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            # 清理单元格文本
                            cell_text = cell.text.strip().replace('\n', ' ')
                            row_data.append(cell_text if cell_text else ' ')
                        table_data.append(row_data)

                    if table_data:
                        print(f"表格 {idx + 1}: {len(table_data)} 行 x {len(table_data[0])} 列")

                        # 创建 PDF 表格
                        t = Table(table_data, hAlign='LEFT')

                        # 设置表格样式
                        table_style = [
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, -1), chinese_font),  # 使用中文字体
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                            ('TOPPADDING', (0, 0), (-1, -1), 6),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ]

                        t.setStyle(TableStyle(table_style))
                        story.append(t)
                        story.append(Spacer(1, 0.3*inch))

                except Exception as e:
                    print(f"表格 {idx + 1} 处理错误: {e}")
                    # 如果表格处理失败，添加提示
                    story.append(Paragraph(f"[表格 {idx + 1}: 转换失败]", chinese_style))
                    story.append(Spacer(1, 0.2*inch))

            # 处理图片
            print(f"文档包含 {len(doc.inline_shapes)} 个内嵌对象（可能包含图片）")
            if len(doc.inline_shapes) > 0:
                from reportlab.platypus import Image as RLImage
                from io import BytesIO
                from PIL import Image as PILImage

                for idx, shape in enumerate(doc.inline_shapes):
                    try:
                        print(f"处理图片 {idx + 1}")

                        # 提取图片数据 - 通过 _inline 访问嵌入的图片
                        inline = shape._inline
                        graphic = inline.graphic
                        pic = graphic.graphicData.pic
                        blip = pic.blipFill.blip
                        embed = blip.embed

                        # 从文档部分获取图片
                        image_part = doc.part.related_parts[embed]
                        image_bytes = image_part.blob

                        # 打开图片
                        image = PILImage.open(BytesIO(image_bytes))

                        # 计算合适的尺寸（保持比例，不超过页面宽度）
                        max_width = 6 * inch  # 页面宽度减去边距
                        max_height = 8 * inch

                        img_width, img_height = image.size
                        aspect = img_height / float(img_width)

                        if img_width > max_width:
                            img_width = max_width
                            img_height = img_width * aspect

                        if img_height > max_height:
                            img_height = max_height
                            img_width = img_height / aspect

                        # 创建临时图片流
                        img_buffer = BytesIO()
                        image.save(img_buffer, format='PNG')
                        img_buffer.seek(0)

                        # 添加图片到 PDF
                        rl_image = RLImage(img_buffer, width=img_width, height=img_height)
                        story.append(rl_image)
                        story.append(Spacer(1, 0.2*inch))
                        print(f"  图片 {idx + 1} 添加成功 ({img_width:.0f}x{img_height:.0f})")

                    except Exception as e:
                        print(f"处理图片 {idx + 1} 失败: {e}")
                        import traceback
                        traceback.print_exc()
                        # 添加占位符
                        story.append(Paragraph(f"[图片 {idx + 1}: 无法提取]", chinese_style))
                        story.append(Spacer(1, 0.2*inch))

            # 如果没有内容，添加提示
            if not story:
                story.append(Paragraph("文档为空或无法提取内容", chinese_style))
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph(
                    "提示：复杂文档建议安装 LibreOffice 进行转换以获得最佳效果。",
                    chinese_style
                ))

            # 生成 PDF
            print("生成 PDF...")
            pdf.build(story)
            print("PDF 生成成功")
            return True

        except Exception as e:
            print(f"备用转换方案失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def convert_ofd_to_pdf(self, file_id: int) -> Conversion:
        """
        将 OFD 文档转换为 PDF

        Args:
            file_id: 文件ID

        Returns:
            Conversion: 转换任务记录

        Raises:
            ValueError: 文件不存在或不是 OFD 文件
            Exception: 转换过程中的错误
        """
        # 获取文件记录
        db_file = self.db.query(File).filter(File.id == file_id).first()
        if not db_file:
            raise ValueError(f"文件不存在: {file_id}")

        # 检查文件类型
        file_ext = get_file_extension(db_file.original_name)
        if file_ext not in ['ofd']:
            raise ValueError(f"不支持的 OFD 格式: {file_ext}")

        # 检查源文件是否存在
        if not os.path.exists(db_file.file_path):
            raise ValueError(f"源文件不存在: {db_file.file_path}")

        # 创建转换任务记录
        conversion = Conversion(
            file_id=file_id,
            source_format=file_ext,
            target_format='pdf',
            status='processing'
        )
        self.db.add(conversion)
        self.db.commit()
        self.db.refresh(conversion)

        try:
            # 生成输出文件名和路径
            output_filename = f"{uuid.uuid4()}.pdf"
            output_path = os.path.join(settings.OUTPUT_DIR, output_filename)

            # 使用 PyMuPDF (fitz) 转换 OFD 到 PDF
            success = self._convert_ofd_with_pymupdf(db_file.file_path, output_path)

            if not success:
                raise Exception("OFD 转 PDF 失败：PyMuPDF 转换失败")

            # 更新转换任务状态
            conversion.status = 'completed'
            conversion.result_path = output_path
            conversion.result_filename = output_filename
            conversion.completed_at = datetime.now()

            # 更新原文件状态
            db_file.status = 'converted'

            self.db.commit()
            self.db.refresh(conversion)

            return conversion

        except Exception as e:
            # 转换失败，更新状态
            conversion.status = 'failed'
            conversion.error_message = str(e)
            conversion.completed_at = datetime.now()
            self.db.commit()
            raise Exception(f"OFD 转 PDF 失败: {str(e)}")

    def _convert_ofd_with_pymupdf(self, input_path: str, output_path: str) -> bool:
        """
        使用 PyMuPDF (fitz) 转换 OFD 到 PDF

        Args:
            input_path: 输入 OFD 文件路径
            output_path: 输出 PDF 文件路径

        Returns:
            是否转换成功
        """
        try:
            import fitz  # PyMuPDF

            # 打开 OFD 文件
            print(f"正在打开 OFD 文件: {input_path}")
            ofd_doc = fitz.open(input_path)

            # 创建一个新的 PDF 文档
            print(f"正在转换为 PDF: {output_path}")
            pdf_doc = fitz.open()  # 创建一个空的 PDF 文档

            # 遍历 OFD 的每一页，将其转换为 PDF 页面
            for page_num in range(len(ofd_doc)):
                ofd_page = ofd_doc.load_page(page_num)

                # 将页面渲染为图像
                pix = ofd_page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2倍缩放以提高质量

                # 创建一个新的 PDF 页面
                pdf_page = pdf_doc.new_page(width=ofd_page.rect.width, height=ofd_page.rect.height)

                # 将图像插入到 PDF 页面
                pdf_page.insert_image(pdf_page.rect, pixmap=pix)

            # 保存 PDF 文档
            pdf_doc.save(output_path, garbage=4, deflate=True, clean=True)

            # 关闭文档
            ofd_doc.close()
            pdf_doc.close()

            print("OFD 转 PDF 成功")
            return True

        except ImportError:
            print("PyMuPDF 未安装，尝试使用其他方法")
            # 如果 PyMuPDF 不可用，可以尝试其他库
            return self._convert_ofd_fallback(input_path, output_path)
        except Exception as e:
            print(f"PyMuPDF 转换失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def _convert_ofd_fallback(self, input_path: str, output_path: str) -> bool:
        """
        OFD 转 PDF 备用方案

        尝试使用其他可用的 OFD 转换库

        Args:
            input_path: 输入 OFD 文件路径
            output_path: 输出 PDF 文件路径

        Returns:
            是否转换成功
        """
        try:
            # 尝试使用 ofd2pdf 库（如果安装）
            try:
                from ofd2pdf import OFDReader, OFD2PDF
                print("使用 ofd2pdf 库进行转换")

                reader = OFDReader(input_path)
                converter = OFD2PDF(reader)
                converter.convert(output_path)

                print("ofd2pdf 转换成功")
                return True
            except ImportError:
                print("ofd2pdf 库未安装")
                pass
            except Exception as e:
                print(f"ofd2pdf 转换失败: {e}")
                pass

            # 如果所有方法都失败，返回 False
            print("所有 OFD 转换方法都失败")
            return False

        except Exception as e:
            print(f"OFD 备用转换方案失败: {e}")
            import traceback
            traceback.print_exc()
            return False

    def convert_archive_to_pdf(self, file_id: int) -> Conversion:
        """
        将压缩包内的文件转换为 PDF 并重新打包

        解压压缩包，将支持的文件格式（Word、图片、OFD）转换为 PDF，
        然后将所有转换后的 PDF 打包成新的 zip 文件

        Args:
            file_id: 文件ID

        Returns:
            Conversion: 转换任务记录

        Raises:
            ValueError: 文件不存在或不是压缩包
            Exception: 转换过程中的错误
        """
        import zipfile
        import rarfile
        import tempfile
        import shutil

        # 获取文件记录
        db_file = self.db.query(File).filter(File.id == file_id).first()
        if not db_file:
            raise ValueError(f"文件不存在: {file_id}")

        # 检查文件类型
        file_ext = get_file_extension(db_file.original_name)
        if file_ext not in ['zip', 'rar']:
            raise ValueError(f"不支持的压缩包格式: {file_ext}")

        # 检查源文件是否存在
        if not os.path.exists(db_file.file_path):
            raise ValueError(f"源文件不存在: {db_file.file_path}")

        # 创建转换任务记录
        conversion = Conversion(
            file_id=file_id,
            source_format=file_ext,
            target_format='pdf',
            status='processing'
        )
        self.db.add(conversion)
        self.db.commit()
        self.db.refresh(conversion)

        temp_extract_dir = None
        temp_pdf_dir = None

        try:
            # 创建临时目录
            temp_extract_dir = tempfile.mkdtemp(prefix='extract_')
            temp_pdf_dir = tempfile.mkdtemp(prefix='pdf_')

            print(f"解压压缩包到: {temp_extract_dir}")

            # 解压压缩包
            if file_ext == 'zip':
                with zipfile.ZipFile(db_file.file_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_extract_dir)
            elif file_ext == 'rar':
                with rarfile.RarFile(db_file.file_path, 'r') as rar_ref:
                    rar_ref.extractall(temp_extract_dir)

            # 遍历解压的文件，转换支持的格式
            converted_count = 0
            supported_extensions = ['doc', 'docx', 'png', 'jpg', 'jpeg', 'gif', 'bmp', 'ofd']

            for root, dirs, files in os.walk(temp_extract_dir):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    file_extension = get_file_extension(filename)

                    if file_extension in supported_extensions:
                        try:
                            print(f"转换文件: {filename} ({file_extension})")

                            # 生成输出 PDF 文件名
                            base_name = os.path.splitext(filename)[0]
                            output_pdf = os.path.join(temp_pdf_dir, f"{base_name}.pdf")

                            # 根据文件类型调用相应的转换方法
                            success = False
                            if file_extension in ['doc', 'docx']:
                                success = self._convert_word_direct(file_path, output_pdf)
                            elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
                                success = self._convert_image_direct(file_path, output_pdf)
                            elif file_extension == 'ofd':
                                success = self._convert_ofd_with_pymupdf(file_path, output_pdf)

                            if success:
                                converted_count += 1
                                print(f"  ✓ 转换成功: {filename}")
                            else:
                                print(f"  ✗ 转换失败: {filename}")

                        except Exception as e:
                            print(f"  ✗ 转换失败: {filename}, 错误: {e}")
                            continue

            if converted_count == 0:
                raise Exception("压缩包中没有找到可转换的文件")

            # 将所有转换后的 PDF 打包成 zip
            output_filename = f"{uuid.uuid4()}.zip"
            output_path = os.path.join(settings.OUTPUT_DIR, output_filename)

            print(f"打包 {converted_count} 个 PDF 文件到: {output_path}")

            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in os.walk(temp_pdf_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arcname = os.path.relpath(file_path, temp_pdf_dir)
                        zipf.write(file_path, arcname)

            # 更新转换任务状态
            conversion.status = 'completed'
            conversion.result_path = output_path
            conversion.result_filename = output_filename
            conversion.completed_at = datetime.now()
            conversion.error_message = f"成功转换 {converted_count} 个文件"

            # 更新原文件状态
            db_file.status = 'converted'

            self.db.commit()
            self.db.refresh(conversion)

            print(f"压缩包转 PDF 完成，共转换 {converted_count} 个文件")

            return conversion

        except Exception as e:
            # 转换失败，更新状态
            conversion.status = 'failed'
            conversion.error_message = str(e)
            conversion.completed_at = datetime.now()
            self.db.commit()
            raise Exception(f"压缩包转 PDF 失败: {str(e)}")

        finally:
            # 清理临时目录
            if temp_extract_dir and os.path.exists(temp_extract_dir):
                try:
                    shutil.rmtree(temp_extract_dir)
                except Exception as e:
                    print(f"清理临时目录失败: {e}")

            if temp_pdf_dir and os.path.exists(temp_pdf_dir):
                try:
                    shutil.rmtree(temp_pdf_dir)
                except Exception as e:
                    print(f"清理临时目录失败: {e}")

    def _convert_word_direct(self, input_path: str, output_path: str) -> bool:
        """
        直接转换 Word 文件到 PDF（不通过数据库）

        Args:
            input_path: 输入 Word 文件路径
            output_path: 输出 PDF 文件路径

        Returns:
            是否转换成功
        """
        try:
            # 尝试使用 LibreOffice
            success = self._convert_with_libreoffice(input_path, output_path)
            if success:
                return True

            # LibreOffice 不可用，使用备用方案
            success = self._convert_word_fallback(input_path, output_path)
            return success

        except Exception as e:
            print(f"Word 转换失败: {e}")
            return False

    def _convert_image_direct(self, input_path: str, output_path: str) -> bool:
        """
        直接转换图片到 PDF（不通过数据库）

        Args:
            input_path: 输入图片文件路径
            output_path: 输出 PDF 文件路径

        Returns:
            是否转换成功
        """
        try:
            from PIL import Image

            image = Image.open(input_path)

            # 转换图片模式
            if image.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', image.size, (255, 255, 255))
                if image.mode == 'P':
                    image = image.convert('RGBA')
                background.paste(image, mask=image.split()[-1] if image.mode in ('RGBA', 'LA') else None)
                image = background
            elif image.mode != 'RGB':
                image = image.convert('RGB')

            # 保存为 PDF
            image.save(output_path, 'PDF', resolution=100.0, quality=95)
            return True

        except Exception as e:
            print(f"图片转换失败: {e}")
            return False

    def get_conversion_by_id(self, conversion_id: int) -> Optional[Conversion]:
        """
        根据ID获取转换任务

        Args:
            conversion_id: 转换任务ID

        Returns:
            转换任务记录或 None
        """
        return self.db.query(Conversion).filter(Conversion.id == conversion_id).first()

    def get_conversions_by_file_id(self, file_id: int) -> list[Conversion]:
        """
        获取文件的所有转换任务

        Args:
            file_id: 文件ID

        Returns:
            转换任务列表
        """
        return self.db.query(Conversion).filter(Conversion.file_id == file_id).all()

    def delete_conversion(self, conversion_id: int) -> bool:
        """
        删除转换任务及其结果文件

        Args:
            conversion_id: 转换任务ID

        Returns:
            是否删除成功
        """
        conversion = self.get_conversion_by_id(conversion_id)
        if not conversion:
            return False

        # 删除结果文件
        if conversion.result_path and os.path.exists(conversion.result_path):
            try:
                os.remove(conversion.result_path)
            except Exception as e:
                print(f"删除转换结果文件失败: {e}")

        # 删除数据库记录
        self.db.delete(conversion)
        self.db.commit()

        return True
